"""
Generate a colourful Word document - BR4C Platform & Admin Guide
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Style Helpers ──
def set_cell_bg(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_colored_heading(text, level=1, color=RGBColor(0x00, 0xBF, 0xD8)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def add_para(text, bold=False, color=None, size=11, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, bold_prefix="", color=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        if color:
            r.font.color.rgb = color
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p

def add_info_box(text, bg_color="E8F4FD", border_color="00BFD8"):
    """Add a styled info box using a single-cell table"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run("💡 " + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0x76)
    set_cell_bg(cell, bg_color)
    # Border
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="8" w:color="{border_color}"/>'
        f'  <w:left w:val="single" w:sz="8" w:color="{border_color}"/>'
        f'  <w:bottom w:val="single" w:sz="8" w:color="{border_color}"/>'
        f'  <w:right w:val="single" w:sz="8" w:color="{border_color}"/>'
        f'</w:tcBorders>'
    )
    tc_pr.append(borders)
    doc.add_paragraph()  # spacer

def add_warning_box(text):
    add_info_box(text, bg_color="FFF3CD", border_color="FFC107")

def add_styled_table(headers, rows, header_bg="1E293B", header_fg=RGBColor(0xFF, 0xFF, 0xFF)):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Style header
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = header_fg
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, header_bg)
    # Rows
    for ri, row in enumerate(rows):
        bg = "F8FAFC" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = table.cell(ri + 1, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            set_cell_bg(cell, bg)
    doc.add_paragraph()  # spacer
    return table

# ═══════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("🎱 BREAKING RACKS OF CASH")
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0xBF, 0xD8)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Platform & Admin Panel Guide")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0xA7, 0x8B, 0xFA)

doc.add_paragraph()
add_para("Complete instructions for managing your 8-Ball Pool web application,\nadmin panel, subscriptions, payments, and all platform features.",
         size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x64, 0x74, 0x8B))

doc.add_paragraph()
doc.add_paragraph()

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info_p.add_run("Version 1.0  •  March 2026")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════

add_colored_heading("📋 Table of Contents", level=1, color=RGBColor(0x00, 0xBF, 0xD8))

toc_items = [
    "1. Platform Overview",
    "2. Admin Panel — Getting Started",
    "3. Admin Panel — Dashboard",
    "4. Admin Panel — User Management",
    "5. Admin Panel — Payment Requests",
    "6. Admin Panel — Season Management",
    "7. Admin Panel — Subscriptions & VIP",
    "8. Admin Panel — Transactions",
    "9. Admin Panel — Actions Log",
    "10. Webapp Features — How Players Use the App",
    "11. VIP Subscription System (How It Works)",
    "12. Energy System",
    "13. League System",
    "14. Daily Rewards & Tasks",
    "15. Leaderboard & Seasons",
    "16. Wallet & Payments (EVM)",
    "17. Important Notes & Tips",
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. PLATFORM OVERVIEW
# ═══════════════════════════════════════════════════════════════

add_colored_heading("1. 🌐 Platform Overview", level=1)

add_para("Breaking Racks of Cash (BR4C) is a web-based 8-Ball Pool gaming platform where users can:", size=11)

bullets = [
    "Play 8-Ball Pool (Rack Attack mode — pot as many balls as you can in 90 seconds)",
    "Earn Cash (in-game currency) from games, daily rewards, tapping, and tasks",
    "Convert Cash to Tokens and withdraw them",
    "Compete on weekly, monthly, and seasonal leaderboards",
    "Purchase VIP subscriptions for extra energy, multipliers, and perks",
    "Progress through 11 league tiers (Wood → Elite)",
]
for b in bullets:
    add_bullet(b)

doc.add_paragraph()
add_para("The platform has TWO main parts:", bold=True, size=12, color=RGBColor(0xA7, 0x8B, 0xFA))

add_styled_table(
    ["Component", "URL", "Purpose"],
    [
        ["Player Webapp", "Your deployed webapp URL", "Where players play games, earn, manage wallet"],
        ["Admin Panel", "Your admin panel URL", "Where YOU manage users, payments, seasons, VIP"],
    ]
)

add_info_box("Both the webapp and admin panel connect to the same Firebase/Firestore database. Changes you make in the admin panel are reflected instantly in the player webapp.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. ADMIN PANEL — GETTING STARTED
# ═══════════════════════════════════════════════════════════════

add_colored_heading("2. 🔑 Admin Panel — Getting Started", level=1)

add_colored_heading("How to Log In", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))

add_para("The admin panel uses Firebase Email/Password authentication. Here's how to log in:", size=11)

steps = [
    ("Step 1: ", "Open your admin panel URL in any web browser (Chrome recommended)."),
    ("Step 2: ", "You will see a login screen with 'BR4C Admin' title."),
    ("Step 3: ", "Enter your admin Email address."),
    ("Step 4: ", "Enter your admin Password."),
    ("Step 5: ", "Click 'Sign In'."),
    ("Step 6: ", "If credentials are correct, you'll see 'Welcome, Admin!' and be taken to the Dashboard."),
]
for prefix, text in steps:
    add_bullet(text, bold_prefix=prefix, color=RGBColor(0x00, 0xBF, 0xD8))

doc.add_paragraph()

add_warning_box("⚠️ IMPORTANT: Your admin account must be created in Firebase Console → Authentication → Users → Add User. Use a strong password. This is NOT a regular player account — it's specifically for admin access.")

add_colored_heading("How to Create an Admin Account", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))

admin_steps = [
    ("Step 1: ", "Go to Firebase Console (console.firebase.google.com)."),
    ("Step 2: ", "Select your project."),
    ("Step 3: ", "Go to Authentication → Users tab."),
    ("Step 4: ", "Click 'Add User'."),
    ("Step 5: ", "Enter the admin email and a strong password."),
    ("Step 6: ", "Click 'Add User' to save."),
    ("Step 7: ", "Now you can log in to the admin panel with those credentials."),
]
for prefix, text in admin_steps:
    add_bullet(text, bold_prefix=prefix, color=RGBColor(0x00, 0xBF, 0xD8))

doc.add_paragraph()

add_colored_heading("Sidebar Navigation", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_para("After logging in, you'll see a sidebar on the left with these sections:", size=11)

add_styled_table(
    ["Icon", "Menu Item", "What It Does"],
    [
        ["📊", "Dashboard", "Overview of all platform stats at a glance"],
        ["👥", "Users", "View, search, edit, and delete player accounts"],
        ["💰", "Payments", "Review & approve/reject token withdrawal requests"],
        ["📅", "Seasons", "Create and manage competitive seasons"],
        ["👑", "Subscriptions", "Manage VIP users, grant VIP, view purchase history"],
        ["🔄", "Transactions", "View all in-game transaction history"],
        ["🛡️", "Actions Log", "See log of all admin actions taken"],
    ]
)

add_info_box("On mobile devices, the sidebar is hidden by default. Tap the ☰ hamburger icon at the top to open it.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. ADMIN PANEL — DASHBOARD
# ═══════════════════════════════════════════════════════════════

add_colored_heading("3. 📊 Admin Panel — Dashboard", level=1)

add_para("The Dashboard gives you a quick overview of your entire platform with 6 stat cards:", size=11)

add_styled_table(
    ["Card", "What It Shows", "Why It Matters"],
    [
        ["Total Users", "Total number of registered players", "Track your user growth"],
        ["Pending Payments", "Withdrawal requests waiting for your review", "You need to approve/reject these!"],
        ["Active VIPs", "Number of users with active VIP subscriptions", "Track subscription revenue"],
        ["Total Cash", "Sum of all in-game cash across all players", "Monitor in-game economy"],
        ["Total Tokens", "Sum of all tokens across all players", "Track withdrawable token supply"],
        ["Games Played", "Total number of games played by all users", "Measure engagement"],
    ]
)

add_warning_box("⚠️ Check the 'Pending Payments' number regularly! Players are waiting for you to approve their withdrawal requests. Don't leave them waiting too long.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. ADMIN PANEL — USER MANAGEMENT
# ═══════════════════════════════════════════════════════════════

add_colored_heading("4. 👥 Admin Panel — User Management", level=1)

add_colored_heading("Viewing Users", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_para("The Users page shows a table of all registered players with the following columns:", size=11)

add_styled_table(
    ["Column", "Description"],
    [
        ["User", "Player name + truncated user ID"],
        ["Cash", "Current in-game cash balance (green)"],
        ["Tokens", "Current token balance (cyan)"],
        ["Games", "Number of games played"],
        ["VIP", "VIP tier (1=Bronze, 2=Silver, 3=Gold) or dash if free"],
        ["Energy", "Current game energy remaining"],
        ["Actions", "Edit (pencil) and Delete (trash) buttons"],
    ]
)

add_colored_heading("Searching for a User", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_bullet("Use the search bar at the top to search by User ID or Username.")
add_bullet("Type the search term and click 'Search' or press Enter.")
add_bullet("Clear the search box and click Search to show all users again.")

add_colored_heading("Editing a User", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_para("Click the pencil icon (✏️) next to any user to open the Edit modal. You can change:", size=11)

edit_fields = [
    ("Cash Balance: ", "Add or remove in-game cash from the player's account."),
    ("Token Balance: ", "Adjust the player's withdrawable token amount."),
    ("Energy: ", "Set how many game plays the user has remaining today."),
    ("VIP Tier: ", "Set the VIP level (0=Free, 1=Bronze, 2=Silver, 3=Gold)."),
    ("Cooldown Resets: ", "Number of tap cooldown resets available."),
]
for prefix, text in edit_fields:
    add_bullet(text, bold_prefix=prefix, color=RGBColor(0x00, 0xBF, 0xD8))

doc.add_paragraph()
add_info_box("After making changes, click 'Save'. The changes are applied instantly to the player's account.")

add_colored_heading("Deleting a User", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_para("Click the trash icon (🗑️) to permanently delete a user. A confirmation dialog will appear. This action CANNOT be undone.", size=11,
         color=RGBColor(0xEF, 0x44, 0x44))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. ADMIN PANEL — PAYMENT REQUESTS
# ═══════════════════════════════════════════════════════════════

add_colored_heading("5. 💰 Admin Panel — Payment Requests", level=1)

add_para("When a player converts Cash to Tokens and requests a withdrawal, it appears here.", size=11)

add_colored_heading("How Payment Requests Work", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))

payment_steps = [
    ("1. ", "Player earns Cash by playing games, tapping, completing tasks."),
    ("2. ", "Player converts Cash → Tokens (rate: 10 Cash = 1 Token) in the Wallet page."),
    ("3. ", "Player submits a withdrawal request with their wallet address."),
    ("4. ", "The request appears in YOUR admin panel as 'Pending'."),
    ("5. ", "YOU review it, then Approve or Reject it."),
]
for prefix, text in payment_steps:
    add_bullet(text, bold_prefix=prefix, color=RGBColor(0x00, 0xBF, 0xD8))

add_colored_heading("Reviewing a Payment", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))

review_steps = [
    ("Step 1: ", "Go to Payments page in the admin panel."),
    ("Step 2: ", "Use the filter buttons (All / Pending / Approved / Rejected) to find pending requests."),
    ("Step 3: ", "Click 'Review' on a pending request."),
    ("Step 4: ", "Optionally add an Admin Note (e.g., 'Paid via ETH on 03/09/2026')."),
    ("Step 5: ", "Click ✅ 'Approve' to approve, or ❌ 'Reject' to reject."),
    ("Step 6: ", "If approved, you need to manually send the tokens/crypto to the player's wallet address shown."),
]
for prefix, text in review_steps:
    add_bullet(text, bold_prefix=prefix, color=RGBColor(0x00, 0xBF, 0xD8))

doc.add_paragraph()
add_warning_box("⚠️ Approving a payment request does NOT automatically send crypto. You must manually transfer the amount to the player's wallet address that is displayed on the request.")

add_colored_heading("Payment Request Details", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_para("Each payment request shows:", size=11)
add_bullet("Amount of Tokens the player wants to withdraw")
add_bullet("User ID of the requesting player")
add_bullet("Wallet Address where they want to receive the payment")
add_bullet("Date/Time of the request")
add_bullet("Status badge (Pending / Approved / Rejected)")
add_bullet("Admin Note (if any)")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. ADMIN PANEL — SEASON MANAGEMENT
# ═══════════════════════════════════════════════════════════════

add_colored_heading("6. 📅 Admin Panel — Season Management", level=1)

add_para("Seasons are competitive periods where players compete on leaderboards for prizes.", size=11)

add_colored_heading("Creating a New Season", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))

season_steps = [
    ("Step 1: ", "Click '+ New Season' button."),
    ("Step 2: ", "Enter a Season Name (e.g., 'Season 1', 'Spring Championship')."),
    ("Step 3: ", "Set the Start Date and End Date."),
    ("Step 4: ", "Click 'Create'."),
]
for prefix, text in season_steps:
    add_bullet(text, bold_prefix=prefix, color=RGBColor(0x00, 0xBF, 0xD8))

add_colored_heading("Managing Season Status", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_para("Each season has a status that you can change:", size=11)

add_styled_table(
    ["Status", "Meaning"],
    [
        ["Active", "Currently running — players can earn season cash and compete"],
        ["Upcoming", "Not started yet — visible but doesn't count points"],
        ["Ended", "Season is over — final leaderboard is locked"],
    ]
)

add_para("To change status: Click 'Edit Status' on any season → select new status → click Save.", size=11)

add_info_box("Only ONE season should be 'Active' at a time. The webapp displays the active season's name and countdown timer in the header.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. ADMIN PANEL — SUBSCRIPTIONS & VIP
# ═══════════════════════════════════════════════════════════════

add_colored_heading("7. 👑 Admin Panel — Subscriptions & VIP", level=1)

add_para("This is one of the most important pages. It has 3 tabs:", size=11, bold=True)

add_colored_heading("Tab 1: Active VIPs", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_para("Shows all users who currently have an active VIP subscription:", size=11)
add_bullet("Player name and ID")
add_bullet("VIP Tier (Bronze / Silver / Gold)")
add_bullet("Expiry date")
add_bullet("Remove button (❌) to cancel their VIP immediately")

add_colored_heading("Tab 2: VIP Purchases", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_para("Shows the history of all VIP subscription purchases:", size=11)
add_bullet("User ID, Tier purchased, Dollar amount paid")
add_bullet("Transaction Hash (from their EVM wallet payment)")
add_bullet("Purchase date and status (confirmed/expired)")

add_colored_heading("Tab 3: Energy Purchases", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_para("Shows all extra energy purchases:", size=11)
add_bullet("User ID, number of energy units bought, dollar amount")
add_bullet("Transaction hash and date")

add_colored_heading("Granting VIP Manually", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_para("You can give VIP to any player for free (e.g., as a promotion or reward):", size=11)

grant_steps = [
    ("Step 1: ", "Click '+ Grant VIP' button (top right)."),
    ("Step 2: ", "Enter the player's User ID (you can find it from the Users page)."),
    ("Step 3: ", "Select the VIP Tier (Bronze, Silver, or Gold)."),
    ("Step 4: ", "Set the Duration in days (default: 30 days)."),
    ("Step 5: ", "Click 'Grant'. The player will instantly get VIP benefits."),
]
for prefix, text in grant_steps:
    add_bullet(text, bold_prefix=prefix, color=RGBColor(0x00, 0xBF, 0xD8))

add_colored_heading("Granting Extra Energy", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_para("You can also give extra game energy to any player:", size=11)
add_bullet("Click '⚡ Grant Energy' button.", bold_prefix="", color=RGBColor(0x00, 0xBF, 0xD8))
add_bullet("Enter User ID and the amount of energy to grant.")
add_bullet("Click 'Grant Energy'. The energy is added to their current balance instantly.")

doc.add_paragraph()
add_info_box("Use 'Grant VIP' for promotional offers or contest winners. Use 'Grant Energy' if a player reports a bug and lost energy.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. ADMIN PANEL — TRANSACTIONS
# ═══════════════════════════════════════════════════════════════

add_colored_heading("8. 🔄 Admin Panel — Transactions", level=1)

add_para("The Transactions page shows ALL in-game financial activity across the platform:", size=11)

add_styled_table(
    ["Column", "Description"],
    [
        ["Type", "Credit (green ↓ — money IN) or Debit (red ↑ — money OUT)"],
        ["User", "The player's ID"],
        ["Amount", "How much cash/tokens were involved"],
        ["Description", "What happened (e.g., 'Game reward', 'Daily reward', 'Withdrawal')"],
        ["Date", "When the transaction occurred"],
    ]
)

add_para("This page is read-only — it's a record of everything that happens on the platform. Use it to:", size=11)
add_bullet("Verify player claims about missing rewards")
add_bullet("Track the in-game economy")
add_bullet("Investigate suspicious activity")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 9. ADMIN PANEL — ACTIONS LOG
# ═══════════════════════════════════════════════════════════════

add_colored_heading("9. 🛡️ Admin Panel — Actions Log", level=1)

add_para("Every action you take in the admin panel is logged here for accountability:", size=11)

add_bullet("What action was taken (e.g., 'Updated user', 'Approved payment', 'Granted VIP')")
add_bullet("Details of the action")
add_bullet("Which admin performed it (email)")
add_bullet("Date and time")

add_info_box("This log cannot be edited or deleted. It serves as an audit trail for all admin actions.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 10. WEBAPP FEATURES — HOW PLAYERS USE THE APP
# ═══════════════════════════════════════════════════════════════

add_colored_heading("10. 🎮 Webapp Features — How Players Use the App", level=1)

add_para("Here's what players experience when they open the webapp:", size=11, bold=True)

add_colored_heading("First Visit — Account Creation", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_bullet("When a player opens the webapp for the first time, a unique account is automatically created.")
add_bullet("They get a randomly generated username (e.g., 'SwiftShark42', 'NeonBreaker7').")
add_bullet("They receive a Welcome Bonus of 500 Cash.")
add_bullet("Their account is stored in the browser — same browser always = same account.")

add_colored_heading("Profile & Name Change", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_bullet("Players can click their profile circle (top-left corner) to open their Profile modal.")
add_bullet("They can change their username ONE TIME only.")
add_bullet("A small pink pen icon appears on the profile circle if they haven't changed their name yet.")
add_bullet("Usernames must be unique — no two players can have the same name.")
add_bullet("After changing, the new name appears everywhere: header, leaderboard, etc.")

add_colored_heading("Home Page", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_bullet("Shows the Rack Attack game button (play 8-Ball Pool)")
add_bullet("Shows the tap-to-earn area for earning cash by tapping")
add_bullet("Daily Reward card (click to claim if available)")
add_bullet("Quick access to all other pages via bottom navigation")

add_colored_heading("Playing Games (Rack Attack)", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_bullet("Players tap 'Rack Attack' to start a solo 90-second pool game.")
add_bullet("Goal: Pot as many balls as you can before time runs out.")
add_bullet("Each game costs 1 Energy. When energy hits 0, they can't play until tomorrow or buy more.")
add_bullet("Cash earned = score × VIP multiplier (1x free, 1.05x Bronze, 1.10x Silver, 1.15x Gold).")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 11. VIP SUBSCRIPTION SYSTEM
# ═══════════════════════════════════════════════════════════════

add_colored_heading("11. 💎 VIP Subscription System (How It Works)", level=1)

add_para("The VIP system is one of the main revenue generators for the platform.", size=11, bold=True)

add_colored_heading("VIP Tiers & Benefits", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))

add_styled_table(
    ["Feature", "Free Tier", "🥉 Bronze ($5/mo)", "🥈 Silver ($10/mo)", "🥇 Gold ($20/mo)"],
    [
        ["Energy/Day", "3", "6", "13", "20"],
        ["Tap Damage", "1x", "6x", "11x", "16x"],
        ["Game Multiplier", "1.00x", "1.05x", "1.10x", "1.15x"],
        ["Cooldown Resets", "0", "2", "5", "10"],
        ["Duration", "—", "30 days", "30 days", "30 days"],
    ],
    header_bg="7C3AED"
)

add_colored_heading("How Players Purchase VIP", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))

vip_purchase = [
    ("1. ", "Player goes to the Shop page in the webapp."),
    ("2. ", "They see the 3 VIP cards (Bronze, Silver, Gold) with prices in USD."),
    ("3. ", "They click 'Subscribe' on their desired plan."),
    ("4. ", "Their EVM-compatible wallet (MetaMask, etc.) opens a payment transaction."),
    ("5. ", "They pay the equivalent amount in ETH/BNB/PLS to your payment wallet."),
    ("6. ", "Once confirmed on-chain, their VIP is activated for 30 days."),
    ("7. ", "The purchase appears in your admin panel under Subscriptions → VIP Purchases."),
]
for prefix, text in vip_purchase:
    add_bullet(text, bold_prefix=prefix, color=RGBColor(0x00, 0xBF, 0xD8))

add_colored_heading("Your Payment Wallet", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_para("All VIP and energy payments are sent to this wallet:", size=11)
p = doc.add_paragraph()
run = p.add_run("0x73F82ecf7345De5d4509be3b418818d00b2cBa1C")
run.font.size = Pt(11)
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0xBF, 0xD8)
run.font.name = "Consolas"

add_para("Supported chains: Ethereum (ETH), Binance Smart Chain (BSC), PulseChain (PLS)", size=10,
         color=RGBColor(0x64, 0x74, 0x8B))

add_colored_heading("VIP Expiry", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_para("VIP subscriptions automatically expire after 30 days. When expired:", size=11)
add_bullet("Player is downgraded back to Free tier")
add_bullet("Energy resets to 3/day")
add_bullet("Tap damage returns to 1x")
add_bullet("Game multiplier returns to 1.00x")
add_info_box("Players must re-purchase VIP when it expires. There is no auto-renewal.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 12. ENERGY SYSTEM
# ═══════════════════════════════════════════════════════════════

add_colored_heading("12. ⚡ Energy System", level=1)

add_para("Energy controls how many games a player can play per day.", size=11, bold=True)

add_styled_table(
    ["VIP Tier", "Energy Per Day", "Extra Energy Available"],
    [
        ["Free", "3 games/day", "Yes, buy in Shop"],
        ["Bronze", "6 games/day", "Yes"],
        ["Silver", "13 games/day", "Yes"],
        ["Gold", "20 games/day", "Yes"],
    ]
)

add_colored_heading("Energy Resets", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_bullet("Energy resets automatically at the start of each new day.")
add_bullet("Reset gives the player their full daily energy based on their VIP tier.")
add_bullet("If a player runs out of energy, they're redirected to the Shop to buy more.")

add_colored_heading("Buying Extra Energy", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_para("Players can buy extra energy in the Shop with real crypto:", size=11)

add_styled_table(
    ["Package", "Energy", "Price (USD)"],
    [
        ["Small", "5 games", "$0.50"],
        ["Medium", "10 games", "$1.00"],
        ["Pack", "25 games", "$2.50"],
        ["Large", "50 games", "$5.00"],
        ["Mega", "100 games", "$10.00"],
    ]
)

add_info_box("Energy purchases appear in the admin panel under Subscriptions → Energy Purchases tab.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 13. LEAGUE SYSTEM
# ═══════════════════════════════════════════════════════════════

add_colored_heading("13. 🏆 League System", level=1)

add_para("Players progress through leagues based on their total Season Cash earned:", size=11)

add_styled_table(
    ["League", "Cash Required", "Icon"],
    [
        ["Wood", "0", "🪵"],
        ["Bronze", "5,000", "🥉"],
        ["Silver", "50,000", "🥈"],
        ["Gold", "250,000", "🥇"],
        ["Platinum", "500,000", "💎"],
        ["Diamond", "1,000,000", "💠"],
        ["Master", "5,000,000", "👑"],
        ["Grandmaster", "10,000,000", "⚔️"],
        ["Legendary", "25,000,000", "🌟"],
        ["Mythic", "50,000,000", "🔮"],
        ["Elite", "100,000,000", "🔥"],
    ]
)

add_para("Players can see their current league in the header bar. Tapping it shows the full Leagues page with all tiers.", size=11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 14. DAILY REWARDS & TASKS
# ═══════════════════════════════════════════════════════════════

add_colored_heading("14. 🎁 Daily Rewards & Tasks", level=1)

add_colored_heading("Daily Rewards (7-Day Cycle)", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_para("Players can claim a free cash reward once per day:", size=11)

add_styled_table(
    ["Day", "Reward"],
    [
        ["Day 1", "50 Cash"],
        ["Day 2", "100 Cash"],
        ["Day 3", "150 Cash"],
        ["Day 4", "200 Cash"],
        ["Day 5", "250 Cash"],
        ["Day 6", "350 Cash"],
        ["Day 7", "500 Cash (Best!)"],
    ],
    header_bg="059669"
)

add_para("After Day 7, the cycle resets back to Day 1.", size=11)
add_para("If a player misses a day, they continue from where they left off (they don't lose progress).", size=11)

add_colored_heading("Tasks", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_para("Players can complete tasks to earn bonus cash:", size=11)

add_styled_table(
    ["Task", "Reward"],
    [
        ["Connect Wallet", "3,000 Cash"],
        ["Invite 1 Friend", "1,000 Cash"],
        ["Invite 5 Friends", "5,000 Cash"],
        ["Invite 10 Friends", "10,000 Cash"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 15. LEADERBOARD & SEASONS
# ═══════════════════════════════════════════════════════════════

add_colored_heading("15. 🏅 Leaderboard & Seasons", level=1)

add_para("The leaderboard has 3 tabs:", size=11, bold=True)

add_styled_table(
    ["Tab", "Resets", "Based On"],
    [
        ["Weekly", "Every week", "Weekly cash earned"],
        ["Monthly", "Every month", "Monthly cash earned"],
        ["Season", "When season ends", "Total season cash earned"],
    ]
)

add_colored_heading("Season Prizes (Top 50)", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))

add_styled_table(
    ["Rank", "Prize (Cash)"],
    [
        ["🥇 1st", "1,000"],
        ["🥈 2nd", "500"],
        ["🥉 3rd", "250"],
        ["4th", "200"],
        ["5th", "150"],
        ["6th–7th", "110–120"],
        ["8th–10th", "70–100"],
        ["11th–20th", "60 each"],
        ["21st–30th", "30 each"],
        ["31st–40th", "15 each"],
        ["41st–50th", "5 each"],
    ],
    header_bg="B45309"
)

add_info_box("You control seasons from the admin panel. Make sure to create a new season before the current one ends! Players need an active season to compete.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 16. WALLET & PAYMENTS (EVM)
# ═══════════════════════════════════════════════════════════════

add_colored_heading("16. 💳 Wallet & Payments (EVM)", level=1)

add_colored_heading("How the Wallet Works for Players", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))

add_bullet("Players go to the Wallet page and click 'Connect Wallet'.")
add_bullet("They connect an EVM-compatible wallet (MetaMask, Trust Wallet, etc.).")
add_bullet("Once connected, their wallet address is saved to their account.")
add_bullet("Connected wallet is shown in the header bar (truncated).")

add_colored_heading("Converting Cash to Tokens", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))
add_para("Conversion rate: 10 Cash = 1 Token", size=12, bold=True, color=RGBColor(0x00, 0xBF, 0xD8))

add_bullet("Players can convert their Cash balance into Tokens on the Wallet page.")
add_bullet("Tokens are what they can request to withdraw as real crypto.")

add_colored_heading("Withdrawal Process", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))

withdraw_steps = [
    ("1. ", "Player must have a wallet connected and minimum 100 tokens."),
    ("2. ", "They enter the amount they want to withdraw."),
    ("3. ", "They click 'Withdraw' and submit the request."),
    ("4. ", "The request appears in YOUR admin panel → Payments page."),
    ("5. ", "YOU review it and manually send crypto to their wallet."),
    ("6. ", "YOU click 'Approve' in the admin panel to mark it as paid."),
    ("7. ", "There's a 3-day cooldown between withdrawal requests."),
]
for prefix, text in withdraw_steps:
    add_bullet(text, bold_prefix=prefix, color=RGBColor(0x00, 0xBF, 0xD8))

add_warning_box("⚠️ IMPORTANT: Withdrawals are NOT automatic. You must manually send crypto to the player's wallet and then approve the request in the admin panel.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 17. IMPORTANT NOTES & TIPS
# ═══════════════════════════════════════════════════════════════

add_colored_heading("17. 📝 Important Notes & Tips", level=1)

add_colored_heading("For You (Admin)", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))

tips = [
    "Check 'Pending Payments' daily — don't keep players waiting.",
    "Always have ONE active season running so leaderboards work.",
    "Use 'Grant VIP' and 'Grant Energy' as promotional tools to retain players.",
    "The Actions Log records everything you do — use it for accountability.",
    "All player data is stored in Firebase/Firestore. Regular backups are recommended.",
    "Your payment wallet receives ALL VIP and energy payments directly — no middleman.",
    "Monitor the Dashboard regularly to understand platform growth.",
]
for t in tips:
    add_bullet(t, bold_prefix="✅ ")

add_colored_heading("Firebase Setup Reminders", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))

firebase_tips = [
    "Enable 'Anonymous Authentication' in Firebase Console → Auth → Sign-in method. This allows webapp users to access the database.",
    "Create your admin account in Firebase Console → Authentication → Users → Add User.",
    "Make sure Firestore rules allow read/write for authenticated users.",
    "Your Firebase config (API keys, project ID) should be set in the .env file.",
]
for t in firebase_tips:
    add_bullet(t, bold_prefix="🔧 ")

add_colored_heading("Security Tips", level=2, color=RGBColor(0xA7, 0x8B, 0xFA))

security_tips = [
    "Never share your admin login credentials with unauthorized people.",
    "Use a strong password (12+ characters) for the admin account.",
    "Your EVM payment wallet private key should be kept offline and secure.",
    "Regularly check the Actions Log for any unauthorized admin activity.",
    "Keep your Firebase project settings private.",
]
for t in security_tips:
    add_bullet(t, bold_prefix="🔒 ")

doc.add_paragraph()
doc.add_paragraph()

# Final footer
final = doc.add_paragraph()
final.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = final.add_run("— End of Guide —")
run.font.size = Pt(14)
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0xBF, 0xD8)

footer_note = doc.add_paragraph()
footer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_note.add_run("Breaking Racks of Cash • Platform Guide v1.0 • March 2026")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

# ── Save ──
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "BR4C_Platform_Guide.docx")
doc.save(output_path)
print(f"✅ Guide saved to: {output_path}")
